from pydantic import BaseModel, Field
import os
import re
import fitz
import docx
import openpyxl
import hashlib
from pathlib import Path
from typing import Optional
from collections import Counter
from schemas.common import DocType
from logger import get_logger

logger = get_logger(__name__)


class ParsedDocument(BaseModel):

    filename: str
    file_path: str
    raw_text: str
    page_count: Optional[int] = None
    file_size_kb: float
    doc_type: DocType
    source_name: str
    issuing_authority: Optional[str] = None
    issue_date: Optional[str] = None
    reference_number: Optional[str] = None
    tags: list[str] = Field(default_factory=list)
    chapter: Optional[str] = None
    section: Optional[str] = None
    notification_number: Optional[str] = None
    hs_code: Optional[str] = None
    customs_section: Optional[str] = None
    document_hash: str


def fingerprint(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def clean_extracted_text(text: str) -> str:
    logger.info("clean_extracted_text | start | chars=%d", len(text))

    text = text.replace("\r", "\n")
    # Remove inline footnote number markers like: 2[, 1[(a), 3[Provided
    text = re.sub(r"\b\d{1,2}\[", "[", text)
    # Remove omitted-text markers: 4*  *  *  * or *  *  *
    text = re.sub(r"^\d*\s*\*[\s\*]*$", "", text, flags=re.MULTILINE)
    # Remove standalone page numbers (just a number on its own line)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    # Remove "Page N" lines
    text = re.sub(r"^\s*Page\s+\d+\s*$", "", text, flags=re.MULTILINE | re.IGNORECASE)
    # Remove long separator lines (underscores, dashes)
    text = re.sub(r"[_\-]{5,}", "", text)
    # Remove "Subject to verification" boilerplate
    text = re.sub(r"\*Subject to verification.*", "", text, flags=re.IGNORECASE)
    # Preserve table rows (lines with |) from being joined
    text = re.sub(r"(?<!\n)\n(?!\n)(?!\[PAGE)(?!.*\|)", " ", text)
    # Remove footnote lines: "1. Subs. by Act...", "2. Ins. by...", "1. 2nd August..."
    text = re.sub(
        r"^\d+\.\s+(Subs\.|Ins\.|Omitted|Inserted|Added|Renumbered).*$",
        "",
        text,
        flags=re.MULTILINE | re.IGNORECASE,
    )
    text = re.sub(
        r"^\d+\.\s+\d{1,2}(st|nd|rd|th)\s+\w+.*$",
        "",
        text,
        flags=re.MULTILINE | re.IGNORECASE,
    )

    # Normalize multiple blank lines → single blank line
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    # Normalize multiple spaces/tabs → single space
    text = re.sub(r"[ \t]+", " ", text)

    logger.info("clean_extracted_text | done | chars_after=%d", len(text))
    return text.strip()


def remove_repeated_lines(text: str) -> str:
    pages = re.split(r"\[PAGE\s+\d+\]", text)
    logger.info("remove_repeated_lines | total_pages=%d", len(pages))

    page_lines = []
    counter = Counter()
    for page in pages:
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        page_lines.append(lines)
        for line in set(lines):
            counter[line] += 1

    threshold = max(3, min(int(len(page_lines) * 0.4), 15))
    repeated = {
        line
        for line, count in counter.items()
        if count >= threshold and 5 < len(line) < 80
    }
    logger.info(
        "remove_repeated_lines | threshold=%d | repeated_lines_removed=%d",
        threshold,
        len(repeated),
    )

    cleaned = []
    for page_index, page in enumerate(page_lines, start=1):
        cleaned.append(f"[PAGE {page_index}]")
        filtered = [line for line in page if line not in repeated]
        cleaned.extend(filtered)

    result = "\n".join(cleaned)
    logger.info("remove_repeated_lines | done | chars_after=%d", len(result))
    return result


def validate_extraction(text: str):
    report = {}
    report["chars"] = len(text)
    report["page_markers"] = text.count("[PAGE")
    report["blank_lines"] = text.count("\n\n")
    report["ocr_noise"] = bool(re.search(r"[^\w\s.,:;()%/\-\[\]]{15,}", text))
    logger.info("validate_extraction | report=%s", report)
    return report


def enrich_metadata(raw_text: str, metadata: dict):
    logger.info("enrich_metadata | start")

    chapter = None
    section = None
    notification = None
    hs_code = None
    customs_section = None

    chapter_match = re.search(r"CHAPTER\s+([IVXLC0-9]+)", raw_text, re.IGNORECASE)
    if chapter_match:
        chapter = chapter_match.group(1)
    logger.info("enrich_metadata | chapter=%s", chapter)

    section_match = re.search(r"\bSection\s+(\d+[A-Z]?)", raw_text, re.IGNORECASE)
    if section_match:
        section = section_match.group(1)
    logger.info("enrich_metadata | section=%s", section)

    notification_match = re.search(
        r"Notification\s+No\.?\s*([\w/-]+)", raw_text, re.IGNORECASE
    )
    if notification_match:
        notification = notification_match.group(1)
    logger.info("enrich_metadata | notification_number=%s", notification)

    hs_match = re.search(r"\b\d{4}(?:\.\d{2})?(?:\.\d{2})?\b", raw_text)
    if hs_match:
        candidate = hs_match.group(0)
        digits = candidate.replace(".", "")
        if len(digits) in [4, 6, 8]:
            hs_code = candidate
    logger.info("enrich_metadata | hs_code=%s", hs_code)

    customs_match = re.search(
        r"\bSection\s+(\d+[A-Z]?)\s+of\s+the\s+Customs\s+Act", raw_text, re.IGNORECASE
    )
    if customs_match:
        customs_section = customs_match.group(1)
    logger.info("enrich_metadata | customs_section=%s", customs_section)

    metadata.update(
        {
            "chapter": chapter,
            "section": section,
            "notification_number": notification,
            "hs_code": hs_code,
            "customs_section": customs_section,
        }
    )

    logger.info("enrich_metadata | done | enriched_keys=%s", list(metadata.keys()))
    return metadata


def extract_pdf_page(page):
    result_parts = []

    tabs = page.find_tables()
    table_bboxes = []

    for table in tabs.tables:
        table_bboxes.append(table.bbox)
        rows = table.extract()
        for row in rows:
            cells = [str(cell).strip() if cell else "" for cell in row]
            if any(cells):
                result_parts.append(" | ".join(cells))

    logger.info(
        "extract_pdf_page | page=%s | tables_found=%d",
        page.number + 1,
        len(table_bboxes),
    )

    blocks = page.get_text("blocks")
    blocks = sorted(blocks, key=lambda b: (round(b[1]), round(b[0])))

    skipped_blocks = 0
    for block in blocks:
        bx0, by0, bx1, by1 = block[:4]
        content = block[4].strip()
        if not content:
            continue
        in_table = any(
            bx0 >= tbbox[0] - 2
            and by0 >= tbbox[1] - 2
            and bx1 <= tbbox[2] + 2
            and by1 <= tbbox[3] + 2
            for tbbox in table_bboxes
        )
        if not in_table:
            result_parts.append(content)
        else:
            skipped_blocks += 1

    logger.info(
        "extract_pdf_page | page=%s | total_blocks=%d | skipped_table_blocks=%d | parts_collected=%d",
        page.number + 1,
        len(blocks),
        skipped_blocks,
        len(result_parts),
    )

    return "\n".join(result_parts)


def parse_pdf(file_path: str) -> tuple[str, int]:
    logger.info("parse_pdf | start | file=%s", Path(file_path).name)
    try:
        doc = fitz.open(file_path)
        page_count = len(doc)
        logger.info("parse_pdf | total_pages=%d", page_count)

        pages_text = []
        for page_num, page in enumerate(doc, start=1):
            text = extract_pdf_page(page)
            if text.strip():
                pages_text.append(f"[PAGE {page_num}]\n{text}")
            else:
                logger.info("parse_pdf | page=%d | skipped (empty)", page_num)

        doc.close()
        joined = "\n\n".join(pages_text)
        logger.info(
            "parse_pdf | done | pages_with_content=%d | total_chars=%d",
            len(pages_text),
            len(joined),
        )
        return (joined, page_count)
    except Exception as e:
        logger.info(
            "parse_pdf | failed | file=%s | error=%s", Path(file_path).name, str(e)
        )
        raise


def parse_docx(file_path: str) -> tuple[str, None]:
    logger.info("parse_docx | start | file=%s", Path(file_path).name)
    document = docx.Document(file_path)
    content = []

    for para in document.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())

    table_rows = 0
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                content.append(" | ".join(cells))
                table_rows += 1

    joined = "\n\n".join(content)
    logger.info(
        "parse_docx | done | paragraphs=%d | table_rows=%d | total_chars=%d",
        len(document.paragraphs),
        table_rows,
        len(joined),
    )
    return (joined, None)


def parse_xlsx(file_path: str) -> tuple[str, None]:
    logger.info("parse_xlsx | start | file=%s", Path(file_path).name)
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    all_text = []
    total_rows = 0

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        all_text.append(f"[SHEET: {sheet}]")
        sheet_rows = 0
        for row_id, row in enumerate(ws.iter_rows(values_only=True), start=1):
            values = [str(cell).strip() for cell in row if cell is not None]
            if values:
                all_text.append(f"[ROW {row_id}] " + " | ".join(values))
                sheet_rows += 1
        logger.info("parse_xlsx | sheet=%s | data_rows=%d", sheet, sheet_rows)
        total_rows += sheet_rows

    wb.close()
    joined = "\n".join(all_text)
    logger.info(
        "parse_xlsx | done | sheets=%d | total_rows=%d | total_chars=%d",
        len(wb.sheetnames),
        total_rows,
        len(joined),
    )
    return (joined, None)


def parse_document(file_path: str, metadata: dict) -> ParsedDocument:
    path = Path(file_path)
    logger.info("parse_document | start | file=%s", path.name)

    if not path.exists():
        logger.info("parse_document | file not found | path=%s", file_path)
        raise FileNotFoundError(f"File not found: {file_path}")

    size_kb = os.path.getsize(file_path) / 1024
    ext = path.suffix.lower()
    logger.info("parse_document | size_kb=%.2f | ext=%s", size_kb, ext)

    if ext == ".pdf":
        raw_text, page_count = parse_pdf(file_path)
    elif ext == ".docx":
        raw_text, page_count = parse_docx(file_path)
    elif ext in (".xlsx", ".xls"):
        raw_text, page_count = parse_xlsx(file_path)
    else:
        logger.info("parse_document | unsupported ext=%s", ext)
        raise ValueError(f"Unsupported type: {ext}")

    logger.info("parse_document | raw extraction done | chars=%d", len(raw_text))

    raw_text = clean_extracted_text(raw_text)
    raw_text = remove_repeated_lines(raw_text)
    validation = validate_extraction(raw_text)
    logger.info("parse_document | validation=%s", validation)

    metadata = enrich_metadata(raw_text, metadata)

    doc_hash = fingerprint(raw_text)
    logger.info("parse_document | document_hash=%s", doc_hash)

    parsed = ParsedDocument(
        filename=path.name,
        file_path=str(path),
        raw_text=raw_text,
        page_count=page_count,
        file_size_kb=round(size_kb, 2),
        doc_type=metadata.get("doc_type", DocType.OTHER),
        source_name=metadata.get("source_name", path.stem),
        issuing_authority=metadata.get("issuing_authority"),
        issue_date=metadata.get("issue_date"),
        reference_number=metadata.get("reference_number"),
        tags=metadata.get("tags", []),
        chapter=metadata.get("chapter"),
        section=metadata.get("section"),
        notification_number=metadata.get("notification_number"),
        hs_code=metadata.get("hs_code"),
        customs_section=metadata.get("customs_section"),
        document_hash=doc_hash,
    )

    logger.info(
        "parse_document | done | file=%s | doc_type=%s | page_count=%s | size_kb=%.2f",
        parsed.filename,
        parsed.doc_type,
        parsed.page_count,
        parsed.file_size_kb,
    )
    return parsed
